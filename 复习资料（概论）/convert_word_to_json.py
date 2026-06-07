from pathlib import Path
from docx import Document
import json
import re


TYPE_NAMES = {
    "single_choice": "单项选择题",
    "multiple_choice": "多项选择题",
    "true_false": "判断题",
    "short_answer": "填空/简答题",
    "essay": "综合应用题",
}

QUESTION_RE = re.compile(r"^(\d+)\s*[.．、]\s*(.*)")
ANSWER_RE = re.compile(r"^(\d+)\s*[.．、]?\s*【参考答案】\s*(.*)")
SECTION_RE = re.compile(r"^([一二三四五六七八九十]+)、(.+?题)")
OPTION_MARK_RE = re.compile(r"(^|\s)([A-D])\s*[.．]")


def clean(text):
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def doc_paras(path):
    return [clean(p.text) for p in Document(str(path)).paragraphs if clean(p.text)]


def extract_options(lines):
    text = clean(" ".join(lines))
    marks = list(OPTION_MARK_RE.finditer(text))
    options = {}
    for i, mark in enumerate(marks):
        key = mark.group(2)
        start = mark.end()
        end = marks[i + 1].start() if i + 1 < len(marks) else len(text)
        value = clean(text[start:end])
        if value:
            options[key] = value
    return options


def parse_questions(lines, qtype, chapter, prefix):
    questions = []
    current = None
    for line in lines:
        match = QUESTION_RE.match(line)
        if match and not re.match(r"^\d+\.\d+", line):
            if current:
                current["options"] = extract_options(current.pop("_option_lines", []))
                current["question"] = clean(" ".join(current.pop("_stem_lines")))
                questions.append(current)
            number = int(match.group(1))
            current = {
                "id": f"{prefix}-{qtype}-{number}",
                "chapter": chapter,
                "type": qtype,
                "typeName": TYPE_NAMES[qtype],
                "number": number,
                "_stem_lines": [match.group(2)],
                "_option_lines": [],
                "answer": "",
            }
            continue
        if not current:
            continue
        if qtype in ("single_choice", "multiple_choice") and OPTION_MARK_RE.search(line):
            current["_option_lines"].append(line)
        else:
            current["_stem_lines"].append(line)
    if current:
        current["options"] = extract_options(current.pop("_option_lines", []))
        current["question"] = clean(" ".join(current.pop("_stem_lines")))
        questions.append(current)
    return [q for q in questions if q.get("question")]


def parse_answer_section(paras, start, end):
    answers = {}
    explanations = {}
    current = None
    chunks = []

    def flush():
        if current is not None:
            explanations[current] = clean(" ".join(chunks))

    for line in paras[start:end]:
        match = ANSWER_RE.match(line)
        if match:
            flush()
            current = int(match.group(1))
            answers[current] = clean(match.group(2)).rstrip("。")
            chunks = []
        elif current is not None:
            chunks.append(line.replace("【解析】", "解析："))
    flush()
    return answers, explanations


def apply_answers(questions, answers, explanations=None):
    explanations = explanations or {}
    for question in questions:
        number = question["number"]
        if number in answers:
            answer = answers[number]
            if question["type"] == "true_false":
                answer = "正确" if "正确" in answer else ("错误" if "错误" in answer else answer)
            question["answer"] = answer
            if explanations.get(number):
                question["analysis"] = explanations[number]
    return questions


def find_heading(paras, text):
    return next((i for i, para in enumerate(paras) if text in para), -1)


def build_question_bank(path):
    paras = doc_paras(path)
    idx_choice = find_heading(paras, "1.1.1 选择题")
    idx_fill = find_heading(paras, "1.1.2 填空题")
    idx_tf = find_heading(paras, "1.1.3 判断题")
    idx_app = find_heading(paras, "1.1.5 综合应用题")
    idx_ans = find_heading(paras, "1.2 试题参考答案")
    idx_choice_ans = find_heading(paras, "1.2.1 选择题")
    idx_fill_ans = find_heading(paras, "1.2.2 填空题")
    idx_tf_ans = find_heading(paras, "1.2.3 判断题")
    idx_short_ans = find_heading(paras, "1.2.4 简答题")
    idx_app_ans = find_heading(paras, "1.2.5 综合应用题")
    idx_algo_ans = find_heading(paras, "1.2.6 算法设计题")

    groups = [
        ("single_choice", "试题库 / 选择题", paras[idx_choice + 1 : idx_fill], idx_choice_ans + 1, idx_fill_ans),
        ("short_answer", "试题库 / 填空题", paras[idx_fill + 1 : idx_tf], idx_fill_ans + 1, idx_tf_ans),
        ("true_false", "试题库 / 判断题", paras[idx_tf + 1 : idx_app], idx_tf_ans + 1, idx_short_ans),
        (
            "essay",
            "试题库 / 综合应用题",
            paras[idx_app + 1 : idx_ans],
            idx_app_ans + 1,
            idx_algo_ans if idx_algo_ans > idx_app_ans else len(paras),
        ),
    ]

    questions = []
    chapters = []
    for qtype, chapter, lines, answer_start, answer_end in groups:
        section_questions = parse_questions(lines, qtype, chapter, "os-bank")
        answers, explanations = parse_answer_section(paras, answer_start, answer_end)
        apply_answers(section_questions, answers, explanations)
        questions.extend(section_questions)
        if section_questions:
            chapters.append({"title": chapter, "questions": section_questions})
    return {
        "source": path.name,
        "title": "计算机操作系统（慕课版）试题库",
        "chapters": chapters,
        "questions": questions,
    }


def section_type(title):
    if "单项" in title or "选择" in title:
        return "single_choice"
    if "填空" in title:
        return "short_answer"
    if "判断" in title:
        return "true_false"
    return "essay"


def parse_exam_answer_tables(path):
    answer_tables = []
    for table in Document(str(path)).tables:
        rows = [[clean(cell.text) for cell in row.cells] for row in table.rows]
        if len(rows) < 2 or not rows[0] or not rows[1]:
            continue
        if rows[0][0].replace(" ", "") != "题号" or rows[1][0].replace(" ", "") != "答案":
            continue
        pairs = []
        for number, answer in zip(rows[0][1:], rows[1][1:]):
            digits = re.sub(r"\D", "", number)
            if digits and answer:
                pairs.append((int(digits), clean(answer)))
        answer_tables.append(pairs)
    return answer_tables


def build_exam_bank(path):
    paras = doc_paras(path)
    answer_start = find_heading(paras, "参考答案与评分标准")
    body = paras[: answer_start if answer_start > 0 else len(paras)]
    questions = []
    current_exam = "期末考试卷"
    current_type = None
    current_lines = []

    def flush():
        nonlocal current_lines
        if current_type and current_lines:
            chapter = f"{current_exam} / {TYPE_NAMES[current_type]}"
            questions.extend(parse_questions(current_lines, current_type, chapter, "os-final"))
        current_lines = []

    for line in body:
        if "期末考试卷（第" in line or re.search(r"期末考试卷（?第\d+套", line):
            current_exam = line.replace("《操作系统》", "").strip()
            continue
        match = SECTION_RE.match(line)
        if match:
            flush()
            current_type = section_type(match.group(2))
            continue
        if current_type:
            current_lines.append(line)
    flush()

    seen = {}
    for question in questions:
        base = question["id"]
        seen[base] = seen.get(base, 0) + 1
        if seen[base] > 1:
            question["id"] = f"{base}-{seen[base]}"

    groups = []
    for question in questions:
        if not groups or groups[-1][0] != question["chapter"]:
            groups.append((question["chapter"], []))
        groups[-1][1].append(question)

    answer_tables = parse_exam_answer_tables(path)
    table_index = 0
    for _, section_questions in groups:
        if table_index >= len(answer_tables):
            break
        apply_answers(section_questions, dict(answer_tables[table_index]))
        table_index += 1

    return {
        "source": path.name,
        "title": "计算机操作系统（慕课版）期末考试卷及评分标准",
        "chapters": [{"title": chapter, "questions": qs} for chapter, qs in groups if qs],
        "questions": questions,
    }


def write_json(path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def main():
    paths = sorted(Path(".").glob("*.docx*"), key=lambda p: p.stat().st_size)
    if len(paths) < 2:
        raise SystemExit("没有找到两份 Word 文档")

    exam = build_exam_bank(paths[0])
    bank = build_question_bank(paths[1])
    merged = {
        "source": [bank["source"], exam["source"]],
        "title": "计算机操作系统（慕课版）合并题库",
        "chapters": bank["chapters"] + exam["chapters"],
        "questions": bank["questions"] + exam["questions"],
    }

    write_json(Path("计算机操作系统试题库.json"), bank)
    write_json(Path("计算机操作系统期末考试卷题库.json"), exam)
    write_json(Path("计算机操作系统合并题库.json"), merged)

    print(f"试题库题数: {len(bank['questions'])}")
    print(f"期末卷题数: {len(exam['questions'])}")
    print(f"合并题数: {len(merged['questions'])}")
    print(f"缺少答案题数: {sum(1 for q in merged['questions'] if not q.get('answer'))}")


if __name__ == "__main__":
    main()
